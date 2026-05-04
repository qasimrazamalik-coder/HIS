import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    source: str
    text: str


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        doc = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return LoadedDocument(source=path.name, text=clean_text(text))


def load_folder(folder: Path) -> list[LoadedDocument]:
    documents: list[LoadedDocument] = []
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            loaded = load_document(path)
            if loaded.text:
                documents.append(loaded)
    return documents

